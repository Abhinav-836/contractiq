import aiofiles
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


async def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from various file types"""
    try:
        if file_type == '.txt':
            return await extract_from_txt(file_path)
        elif file_type == '.pdf':
            return await extract_from_pdf(file_path)
        elif file_type in ['.doc', '.docx']:
            return await extract_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return f"Unsupported file type: {file_type}"
    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        return f"Error extracting text: {str(e)}"


async def extract_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = await f.read()
            return content.strip()
    except Exception as e:
        logger.error(f"TXT extraction error: {str(e)}")
        return ""


async def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        import PyPDF2
        text_parts = []
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            # Extract metadata
            if reader.metadata:
                metadata = []
                if reader.metadata.title:
                    metadata.append(f"Title: {reader.metadata.title}")
                if reader.metadata.author:
                    metadata.append(f"Author: {reader.metadata.author}")
                if reader.metadata.subject:
                    metadata.append(f"Subject: {reader.metadata.subject}")
                if metadata:
                    text_parts.append("\n".join(metadata))
                    text_parts.append("")
            
            # Extract text from each page
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- Page {i+1} ---")
                    text_parts.append(page_text.strip())
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return f"[PDF extraction error: {e}]"


async def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return f"[DOCX extraction error: {e}]"